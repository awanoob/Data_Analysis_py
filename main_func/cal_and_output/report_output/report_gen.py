import logging
from datetime import datetime
from docx import Document
from docx.enum.section import WD_SECTION, WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
import pandas as pd
import os
import glob
import sys
# from main_func.cal_and_output.map_pic_gen.map_pic import map_generator
# import main_func.cal_and_output.report_output.styles as styles
from cal_and_output.map_pic_gen.map_pic import map_generator
import cal_and_output.report_output.styles as styles

# 打开模板文档
doc = Document(r'.\cal_and_output\report_output\default\module_default.docx')
# doc = Document(r'.\default\module_default.docx')

default_section = doc.sections[0]

# 设置样式
styles.modify_styles(doc)
styles.create_custom_styles(doc)


def get_png_files(folder_path):
    """以每个文件夹名为key，文件夹下的所有.png文件名为value，构建字典"""
    # 获取该文件夹下的所有文件夹的名称
    root_dirs = os.listdir(folder_path)
    png_files = {}
    # 遍历每个根目录
    for root_dir in root_dirs:
        # 遍历根目录及其子目录
        for foldername, subfolders, filenames in os.walk(folder_path + '//' + root_dir):
            # 找到当前目录中的所有 .png 文件
            png_filepaths = glob.glob(os.path.join(foldername, '*.png'))
            if png_filepaths:  # 如果找到 .png 文件
                # 获取当前文件夹的名称作为键
                folder_key = os.path.basename(foldername)
                # 获取文件名作为值
                png_filenames = [os.path.abspath(png) for png in png_filepaths]
                # 将文件夹名与文件名列表添加到字典中
                png_files[folder_key] = png_filenames

    return png_files


def get_dev_datas(folder_path, multi_folder_path):
    """获取csv误差文件并生成地图"""
    # 获取该文件夹下的所有子文件夹的名称
    folder_paths = os.listdir(folder_path)
    err_paths = []
    scene_paths = {}

    # 存储每个场景对应的输出文件夹路径
    scene_output_paths = {}
    scene_folder_paths = []

    for folder in folder_paths:
        # 拼接误差统计表的路径
        err_file = os.path.join(folder_path, folder, f"{folder}_误差统计表.csv")
        err_paths.append(err_file)
        # 获取文件夹下子文件夹的名称（不包含'.'的文件夹）
        for scene_folder in os.listdir(os.path.join(folder_path, folder)):
            if os.path.isdir(os.path.join(folder_path, folder, scene_folder)) and '.' not in scene_folder:
                scene_folder_paths.append(scene_folder)

        # 记录每个场景的输出路径
        for scene_folder in scene_folder_paths:
            scene_output_path = os.path.join(folder_path, scene_folder)
            scene_output_paths[scene_folder] = scene_output_path

            # 初始化 scene_paths 中的 scene_folder 列表
            if scene_folder not in scene_paths:
                scene_paths[scene_folder] = []

            # 拼接 test.navplot 文件路径并添加到 scene_paths 字典的对应列表中
            navplot_file = os.path.join(folder_path, folder, scene_folder, f"{scene_folder}_test.navplot")
            scene_paths[scene_folder].append(navplot_file)

    # 使用字典存储每个文件夹名称与其对应的DataFrame
    xlsx_datas = {}
    for filepath in err_paths:
        folder_name = os.path.basename(os.path.dirname(filepath))
        xlsx_data = pd.read_csv(filepath)

        xlsx_datas[folder_name] = {}

        for i in range(0, len(xlsx_data), 6):
            scene_name = xlsx_data.iloc[i, 0]
            xlsx_datas[folder_name][scene_name] = {}
            xlsx_datas[folder_name][scene_name]['data'] = xlsx_data.iloc[i:i + 5, 2:]
            xlsx_datas[folder_name][scene_name]['fix_ratio'] = float(xlsx_data.iloc[i, -2])
            xlsx_datas[folder_name][scene_name]['distance'] = float(xlsx_data.iloc[i, -1])

    # 为每个场景生成地图
    map_paths = {}
    for scene_name, navplot_files in scene_paths.items():
        try:
            output_path_full = os.path.join(multi_folder_path, scene_name, f"{scene_name}_map_full.png")
            output_path_zoomed = os.path.join(multi_folder_path, scene_name, f"{scene_name}_map_zoomed.png")

            # 调用map_gen生成地图
            map_generator(navplot_files, output_path_full, output_path_zoomed)

            map_paths[scene_name] = (output_path_full, output_path_zoomed)

        except Exception as e:
            logging.error(f"生成{scene_name}场景地图失败: {str(e)}")
            continue

    return xlsx_datas, map_paths


def header_info():
    """在页眉中设置信息"""
    # 获取页眉中的第一个表格
    header_table = doc.sections[0].header.tables[0]
    # 今天的日期
    today = datetime.today().strftime('%m%d-%Y')
    # 设置表格中的内容
    header_table.cell(0, 2).text = 'CHC-JSBG-' + today
    header_table.cell(0, 2).paragraphs[0].style = doc.styles['Header']


def footer_pages():
    """在页脚中设置页码"""
    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    # 获取第二节的页脚
    footer = section2.footer
    footer.is_linked_to_previous = False  # 不使用上一章节的页脚样式
    # 在页脚添加一个段落
    paragraph = footer.add_paragraph(style='Footer')
    # 添加页码相关的字段
    run1 = paragraph.add_run('')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._element.append(fldChar1)

    run2 = paragraph.add_run('')
    fldChar2 = OxmlElement('w:instrText')
    fldChar2.text = 'PAGE'
    run2._element.append(fldChar2)

    run3 = paragraph.add_run('')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'separate')
    run3._element.append(fldChar3)

    run4 = paragraph.add_run('')
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run4._element.append(fldChar4)


def set_cell_border(cell, **kwargs):
    """
    设置表格单元格边框
    :param cell: 单元格对象
    :param kwargs: 例如 top="2.25pt", bottom="2.25pt", start="2.25pt", end="2.25pt" 等
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for key in kwargs:
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        border_attr = OxmlElement(f'w:{key}')
        border_attr.set(qn('w:val'), 'single')  # 单线
        border_attr.set(qn('w:sz'), '4')  # 边框宽度，单位为1/8 pt
        border_attr.set(qn('w:space'), '0')  # 间距
        border_attr.set(qn('w:color'), 'auto')  # 自动颜色

        tcBorders.append(border_attr)


def set_cell_margins(cell, **kwargs):
    '''设置某单元格间距

    长度单位为Twips，1Twips = 1/20pt，1Twips = 1/567cm

    :param cell: 某单元格
    :param top: 上边距
    :param start: 左边距
    :param bottom: 下边距
    :param end: 右边距
    :param left: 左边距（WPS）
    :param right: 右边距（WPS）
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'start', 'bottom', 'end', 'left', 'right']:
        if m in kwargs:
            node = OxmlElement('w:{}'.format(m))
            node.set(qn('w:w'), str(kwargs.get(m) * 567))  # 1cm = 567 Twips
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)


def apply_table_style(table, border_size="2.25pt"):
    """
    为表格中的每个单元格应用全框线样式
    :param table: 表格对象
    :param border_size: 边框宽度
    """
    column_widths = {
        0: 0.8,
        1: 1.1,
        2: 1.0,
        3: 1.3,
        4: 1.3,
        5: 1.3,
        6: 1.3,
        7: 1.6,
        8: 1.6,
        9: 1.6,
        10: 1.6,
        11: 1.4,
        12: 1.4,
        13: 1.4
    }
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.height = Cm(1.3)
        for idx, cell in enumerate(row.cells):
            set_cell_border(cell, top=border_size, bottom=border_size, start=border_size, end=border_size)
            set_cell_margins(cell, start=0.05, end=0.05)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.width = Cm(column_widths.get(idx, 1))


def fill_table_data(table, title, product_names, xlsx_datas):
    """填充表格数据"""
    indicators = ['1σ', '2σ', '3σ', 'RMS', 'MAX']

    for i, product_name in enumerate(product_names):
        device_data = xlsx_datas[product_name][title]['data']

        for j in range(5):
            # 填充指标列
            table.cell(5 * i + j + 1, 2).text = indicators[j]
            table.cell(5 * i + j + 1, 2).paragraphs[0].style = doc.styles['TableParagraph']

            # 填充数据列 - 修改这里,只取前10列数据(不包括固定率和总里程)
            for k in range(10):  # 修改这里,原来是range(12)
                cell_value = device_data.iloc[j, k]  # 修改这里,去掉+2
                formatted_value = f"{cell_value:.4f}" if isinstance(cell_value, (int, float)) else str(cell_value)
                table.cell(5 * i + j + 1, k + 3).text = formatted_value
                table.cell(5 * i + j + 1, k + 3).paragraphs[0].style = doc.styles['TableParagraph']


def write_cover():
    """封面"""
    # 拿到第一个节的第一个表格
    table = doc.tables[0]
    # 今天的日期
    today = datetime.today().strftime('%m%d-%Y')
    # 设置表格中的内容
    table.cell(0, 1).text = ' CHC-JSBG-' + today + '  '
    table.cell(0, 1).paragraphs[0].style = doc.styles['Cover_Table']
    # 拿到第14个段落
    paragraph = doc.paragraphs[13]
    # 设置段落中的内容
    paragraph.text = datetime.today().strftime('%Y年%m月%d日')
    paragraph.style = doc.styles['Cover_Time']


def write_chapter1():
    """第一章的内容"""
    doc.add_paragraph('测试方案概述', style='Heading 1')
    doc.add_paragraph('测试目的', style='Heading 2')
    doc.add_paragraph('')
    doc.add_paragraph('测试日程', style='Heading 2')
    doc.add_paragraph('')
    doc.add_paragraph('测试环境', style='Heading 2')
    doc.add_paragraph('连接框图', style='Heading 3')
    doc.add_paragraph('')
    doc.add_paragraph('测试车辆', style='Heading 3')
    doc.add_paragraph('')
    doc.add_paragraph('测试天线', style='Heading 3')
    doc.add_paragraph('')


def write_chapter2(pic_filepath, dev_filepath):
    """第二章的内容"""
    doc.add_paragraph('结果展示与总结', style='Heading 1')
    # 获取图片文件
    pic_files = get_png_files(pic_filepath)
    # 获取误差文件和地图文件
    xlsx_datas, map_paths = get_dev_datas(dev_filepath, pic_filepath)
    product_names = list(xlsx_datas.keys())

    count = 1
    for title, pics in pic_files.items():
        doc.add_paragraph(title, style='Heading 2')
        doc.add_paragraph('（1）轨迹对比图')
        paragraph = doc.add_paragraph(style='Picture')
        run = paragraph.add_run()
        run.add_picture(map_paths[title][0], width=Cm(15))
        doc.add_paragraph(f'图2.{count} 场景轨迹图', style='PictureName')
        count += 1

        # 处理误差序列图
        doc.add_paragraph('（2）误差序列图')
        paragraph = doc.add_paragraph(style='Picture')
        run = paragraph.add_run()
        run.add_picture(map_paths[title][1], width=Cm(15))
        doc.add_paragraph(f'图2.{count} 局部对比图', style='PictureName')
        count += 1
        # 定义图片类型和顺序
        pic_types = {
            '_xy.png': {'order': 1, 'description': '位置误差-1'},
            '_pos_alt.png': {'order': 2, 'description': '位置误差-2'},
            '_vel.png': {'order': 3, 'description': '速度误差'},
            '_att.png': {'order': 4, 'description': '姿态误差'}
        }

        # 收集和排序误差图片
        error_pics = []
        for pic in pics:
            if 'map' not in pic:  # 排除地图图片
                # 获取图片的类型后缀
                pic_suffix = None
                for suffix in pic_types.keys():
                    if pic.endswith(suffix):
                        pic_suffix = suffix
                        break

                if pic_suffix:
                    error_pics.append({
                        'path': pic,
                        'order': pic_types[pic_suffix]['order'],
                        'description': pic_types[pic_suffix]['description']
                    })

        # 按预定顺序排序
        error_pics.sort(key=lambda x: x['order'])

        # 添加排序后的图片
        for pic_info in error_pics:
            try:
                paragraph = doc.add_paragraph(style='Picture')
                run = paragraph.add_run()
                run.add_picture(pic_info['path'], width=Cm(15))
                doc.add_paragraph(f'图2.{count} {pic_info["description"]}', style='PictureName')
                count += 1
            except Exception as e:
                logging.warning(f"添加图片到文档时出错: {str(e)}")
                continue

        # 处理产品误差统计表格
        doc.add_paragraph('（3）产品误差统计')
        # 添加表格
        rows = len(xlsx_datas) * 5 + 1
        cols = 13
        table = doc.add_table(rows=rows, cols=cols)
        apply_table_style(table)

        # 设置表头
        first_device_data = xlsx_datas[product_names[0]][title]['data']
        table_head = ['场景', '产品', '参数'] + list(first_device_data.columns[0:10])

        # 合并第一列（除表头外的所有单元格）
        first_col_cells = table.column_cells(0)[1:]
        merged_first_col = first_col_cells[0].merge(first_col_cells[-1])
        merged_first_col.text = title
        merged_first_col.paragraphs[0].style = doc.styles['TableParagraph']

        # 合并第二列（除表头外每5个单元格合并一次）
        second_col_cells = table.column_cells(1)[1:]
        for i in range(0, len(second_col_cells), 5):
            cells_to_merge = second_col_cells[i:i + 5]
            merged_cell = cells_to_merge[0].merge(cells_to_merge[-1])
            merged_cell.text = product_names[i // 5]
            merged_cell.paragraphs[0].style = doc.styles['TableParagraph']

        # 设置表头
        for i in range(cols):
            table.cell(0, i).text = table_head[i]
            table.cell(0, i).paragraphs[0].style = doc.styles['TableParagraph']

        # 填充数据
        fill_table_data(table, title, product_names, xlsx_datas)

        # 添加结论段落
        doc.add_paragraph(f'统计结果表明，在{title}道路场景下，全程：{xlsx_datas[product_names[0]][title]["distance"]}米。')

        # 添加设备固定率和距离信息
        for product_name in product_names:
            fix_ratio = xlsx_datas[product_name][title]['fix_ratio']
            doc.add_paragraph(f'{product_name}设备固定率为{fix_ratio:.2f}%。')

        doc.add_paragraph('')


def write_chapter3():
    """第三章的内容"""
    doc.add_paragraph('结果展示与总结', style='Heading 1')
    doc.add_paragraph('水平位置误差RMS汇总', style='Heading 2')
    doc.add_paragraph('')
    doc.add_paragraph('航向角误差RMS汇总', style='Heading 2')


def report_gen_func(input_cfg):
    pic_folder_path = input_cfg['multi_dev_err_path']
    dev_folder_path = input_cfg['path_proj_dev']
    # 读取excel文件

    # 设置页脚信息
    header_info()
    # 封面
    write_cover()

    # 在页脚中设置页码
    footer_pages()

    # 第一章
    write_chapter1()
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    # 第二章
    write_chapter2(pic_folder_path, dev_folder_path)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    # 第三章
    write_chapter3()
    # 保存文档
    doc.save(os.path.join(input_cfg['path_proj'], '测试报告.docx'))


if __name__ == '__main__':
    input_cfg = {
        'path_proj': r'D:\python\zhongqichuangzhi',
        'multi_dev_err_path': r"D:\python\zhongqichuangzhi\multi_dev_err_plot",
        'path_proj_dev': r"D:\python\zhongqichuangzhi\result_all"
    }
    report_gen_func(input_cfg)
    sys.exit(0)
